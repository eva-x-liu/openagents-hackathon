"""
BookClub AI Agent - 多智能体读书会全案生成系统

模块：base_agent.py
描述：BookClub Agent 核心基类，实现 3 个 Agent（Intake/Content/Ops）的共享逻辑

核心功能：
    1. 多 Agent 协作：Intake → Content → Ops 流水线
    2. 知识库集成：PDF + 营养速查表 + 膳食指南
    3. 分天生成策略：避免长文截断，支持任意天数
    4. 多格式输出：Markdown + Word + 微信纯文本

技术亮点：
    - 三层知识库架构（数据调用优先级）
    - 动态 max_output_tokens 设置
    - 渐进式销售策略（Day1种草 → 中间见证 → 最后销讲）

依赖：
    - OpenAgents: 多智能体框架
    - Gemini 2.0 Flash: LLM 引擎
    - python-docx: Word 文档生成

作者：Eva（注册营养师 + INTJ 架构师）
版本：1.0
日期：2026-01-07
许可证：MIT
"""

import os
import sys
import re
from datetime import datetime
from dotenv import load_dotenv
from google import genai

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("⚠️ python-docx 未安装，Word 输出功能不可用", flush=True)

from openagents.agents.worker_agent import (
    WorkerAgent as Agent,
    ChannelMessageContext,
    EventContext,
    on_event,
)

load_dotenv()
sys.path.append(os.getcwd())

# 输出目录（自动保存 Agent 生成的内容）
OUTPUT_DIR = "output"


EVENT_TO_CONTENT = "bookclub.pipeline.to_content"
EVENT_TO_OPS = "bookclub.pipeline.to_ops"


class BookClubAgent(Agent):
    def __init__(self, *args, **kwargs):
        # 先调用父类初始化
        super().__init__(*args, **kwargs)

        # OpenAgents 使用 agent_config 和 agent_id（不是 config 和 id）
        agent_id = kwargs.get('agent_id') or getattr(self, '_agent_id', 'unknown')
        agent_config = kwargs.get('agent_config') or getattr(self, '_agent_config', {})
        
        # 获取 config 部分（agent_config 可能包含多个部分）
        if isinstance(agent_config, dict):
            self.raw_config = agent_config.get('config', agent_config)
        else:
            self.raw_config = {}

        # 优先从配置文件读取 role_type，否则从 agent_id 推断
        self.role_type = self.raw_config.get("role_type")
        
        if not self.role_type:
            agent_id_str = str(agent_id).lower()
            if "content" in agent_id_str:
                self.role_type = "content"
            elif "intake" in agent_id_str:
                self.role_type = "intake"
            else:
                self.role_type = "ops"

        # 从配置读取模型名，支持不同 Agent 使用不同模型
        self.model_name = self.raw_config.get("model_name", "gemini-2.0-flash-exp")
        self.instruction = self.raw_config.get("instruction", "你是一位专业的临床营养专家助手。")

        api_key = os.getenv("GOOGLE_API_KEY")
        self.genai_client = genai.Client(
            api_key=api_key,
            http_options={"api_version": "v1beta"},
        ) if api_key else None

        self.file_ref = None
        self.rules_content = None  # 膳食指南规则（Markdown 文本）
        self.nutrition_content = None  # 食物营养速查表（Markdown 文本）
        print(f"✅ [Ready] {self.role_type.upper()} 就绪 | 引擎: {self.model_name}", flush=True)

    async def on_startup(self):
        """
        Agent 启动时加载知识库并发送欢迎消息
        - content/ops：加载膳食规则
        - content：额外加载 PDF + 营养速查表
        - intake：发送使用指南到频道
        
        【数据调用优先级】
        - 涉及具体克数（g/ml）时 → 优先检索 nutrition_reference.md
        - 涉及医学逻辑时 → 优先检索 you_are_what_you_eat.pdf
        - 涉及定量标准时 → 必须核对 dietary_rules.md
        """
        # content 和 ops 都需要膳食规则
        if self.role_type in ("content", "ops"):
            await self._load_dietary_rules()
        
        # content 需要 PDF 知识库 + 营养速查表
        if self.role_type == "content":
            await self._setup_knowledge_base()
            await self._load_nutrition_data()
        
        # intake 发送欢迎消息到 #general 频道
        if self.role_type == "intake":
            await self._send_welcome_message()
    
    async def _send_welcome_message(self):
        """
        向 #general 频道发送欢迎消息和使用指南
        仅 Intake Agent 在启动时调用一次
        """
        welcome_msg = """
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
📖 欢迎使用 BookClub Core - 读书会全案生成系统
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

🎯 系统说明
基于《你是你吃出来的》，自动生成完整读书会策划方案。
支持 3 种格式输出：Markdown / Word / 微信友好版

🚀 使用流程（3步）

Step 1: @bc-intake → 告诉我你的需求
Step 2: 引用 Intake 消息 → @bc-content → 生成讲书内容
Step 3: 引用 Content 消息 → @bc-ops → 生成执行物料

💡 提示：使用"引用"功能，无需手动复制粘贴！

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

📝 Prompt 模板

@bc-intake

我是注册营养师Eva，想做一个 3 天的读书会。

【书籍信息】
书名：《你是你吃出来的》

【主理人信息】
专业特长：神经营养学 / 慢病管理 / 肠道健康
风格偏好：专业严谨 / 温情亲切 / 幽默风趣

【项目参数】
交付周期：3天 / 5天 / 7天
招募周期：7天

【产品信息】（可选）
产品名：复合B族维生素
功效：支持神经递质、改善疲劳
定位：辅助工具（不是主角）

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

🎯 极简示例

@bc-intake

我是注册营养师Eva，3天读书会
特长：神经营养学
产品：复合B族维生素

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

💡 提示
- 所有输出会自动保存到 output/ 目录（3种格式）
- Word 格式可直接复制到微信公众号编辑器
- 微信版文本适合朋友圈逐条复制

🔗 详细文档：使用指南.md

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
✅ 系统已就绪，随时开始！
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
"""
        
        try:
            ws = self.workspace()
            # 发送到 #general 频道
            await ws.channel("general").post(welcome_msg)
            print(f"✅ [Welcome] 欢迎消息已发送到 #general 频道", flush=True)
        except Exception as e:
            print(f"⚠️ [Welcome] 欢迎消息发送失败: {e}", flush=True)
    
    async def _load_dietary_rules(self):
        """
        加载膳食指南规则（Markdown 文本）
        规则文件：data/dietary_rules.md
        """
        rules_path = "data/dietary_rules.md"
        if not os.path.exists(rules_path):
            print(f"⚠️ [System] 膳食规则文件不存在: {rules_path}", flush=True)
            return
        
        try:
            with open(rules_path, "r", encoding="utf-8") as f:
                self.rules_content = f.read()
            print(f"✅ [System] 膳食规则已加载（{len(self.rules_content)} 字符）", flush=True)
        except Exception as e:
            print(f"💥 [System] 膳食规则加载失败: {e}", flush=True)

    async def _load_nutrition_data(self):
        """
        加载食物营养速查表（Markdown 文本）
        规则文件：data/nutrition_reference.md
        
        【数据调用优先级】
        - 涉及具体克数（g/ml）时 → 优先检索此表
        - 涉及医学逻辑时 → 优先检索 PDF
        - 涉及定量标准时 → 必须核对 dietary_rules.md
        """
        nutrition_path = "data/nutrition_reference.md"
        if not os.path.exists(nutrition_path):
            print(f"⚠️ [System] 营养速查表不存在: {nutrition_path}", flush=True)
            return
        
        try:
            with open(nutrition_path, "r", encoding="utf-8") as f:
                self.nutrition_content = f.read()
            print(f"✅ [System] 营养速查表已加载（{len(self.nutrition_content)} 字符）", flush=True)
        except Exception as e:
            print(f"💥 [System] 营养速查表加载失败: {e}", flush=True)

    def _save_output(self, content: str, suffix: str = "") -> str:
        """
        自动保存 Agent 输出到多种格式
        - .md：Markdown 原文（开发者查看）
        - .docx：Word 文档（微信公众号编辑器）
        - _wechat.txt：微信友好纯文本（朋友圈复制）
        
        Args:
            content: 要保存的内容（Markdown 格式）
            suffix: 可选后缀（如 "day1"）
        
        Returns:
            Markdown 文件路径（作为主路径）
        """
        # 确保 output 目录存在
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        
        # 生成基础文件名：角色_日期时间_后缀
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"{self.role_type}_{timestamp}"
        if suffix:
            base_filename += f"_{suffix}"
        
        saved_files = []
        
        # 1. 保存 Markdown 原文
        md_filepath = os.path.join(OUTPUT_DIR, f"{base_filename}.md")
        try:
            with open(md_filepath, "w", encoding="utf-8") as f:
                f.write(content)
            saved_files.append(f"{base_filename}.md")
        except Exception as e:
            print(f"💥 [Save] Markdown 保存失败: {e}", flush=True)
        
        # 2. 生成 Word 文档
        if DOCX_AVAILABLE:
            docx_filepath = os.path.join(OUTPUT_DIR, f"{base_filename}.docx")
            try:
                self._markdown_to_docx(content, docx_filepath)
                saved_files.append(f"{base_filename}.docx")
            except Exception as e:
                print(f"💥 [Save] Word 转换失败: {e}", flush=True)
        
        # 3. 生成微信友好版
        wechat_filepath = os.path.join(OUTPUT_DIR, f"{base_filename}_wechat.txt")
        try:
            wechat_content = self._markdown_to_wechat(content)
            with open(wechat_filepath, "w", encoding="utf-8") as f:
                f.write(wechat_content)
            saved_files.append(f"{base_filename}_wechat.txt")
        except Exception as e:
            print(f"💥 [Save] 微信版转换失败: {e}", flush=True)
        
        # 输出保存结果
        if saved_files:
            print(f"💾 [Save] 已生成 {len(saved_files)} 个文件:", flush=True)
            for f in saved_files:
                print(f"  - output/{f}", flush=True)
        
        return md_filepath
    
    def _markdown_to_docx(self, markdown_text: str, output_path: str):
        """
        将 Markdown 转换为 Word 文档（微信公众号编辑器友好）
        
        支持：
        - # 标题（一级到六级）
        - **粗体**
        - - 列表
        - | 表格 |
        - 分隔线
        """
        doc = Document()
        
        # 设置默认字体（微软雅黑，微信编辑器友好）
        style = doc.styles['Normal']
        style.font.name = '微软雅黑'
        style.font.size = Pt(12)
        
        lines = markdown_text.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # 跳过空行
            if not line:
                i += 1
                continue
            
            # 一级标题 (# )
            if line.startswith('# ') and not line.startswith('## '):
                text = line[2:].strip()
                p = doc.add_heading(text, level=1)
                p.runs[0].font.name = '微软雅黑'
                p.runs[0].font.bold = True
            
            # 二级标题 (## )
            elif line.startswith('## ') and not line.startswith('### '):
                text = line[3:].strip()
                p = doc.add_heading(text, level=2)
                p.runs[0].font.name = '微软雅黑'
            
            # 三级标题 (### )
            elif line.startswith('### '):
                text = line[4:].strip()
                p = doc.add_heading(text, level=3)
                p.runs[0].font.name = '微软雅黑'
            
            # 列表项 (- 或 * 开头)
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                # 处理粗体 **text**
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                p = doc.add_paragraph(text, style='List Bullet')
            
            # 数字列表 (1. 开头)
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line).strip()
                text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
                p = doc.add_paragraph(text, style='List Number')
            
            # 表格（简单检测 | 开头）
            elif line.startswith('|'):
                # 收集连续的表格行
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i].strip())
                    i += 1
                i -= 1  # 回退一行
                
                # 过滤掉分隔行 (|---|---|)
                table_lines = [l for l in table_lines if not re.match(r'^\|[\s\-:]+\|', l)]
                
                if table_lines:
                    # 解析表格
                    rows = []
                    for tl in table_lines:
                        cells = [c.strip() for c in tl.split('|')[1:-1]]  # 去掉首尾空元素
                        rows.append(cells)
                    
                    if rows:
                        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                        table.style = 'Light Grid Accent 1'
                        
                        for row_idx, row_data in enumerate(rows):
                            for col_idx, cell_text in enumerate(row_data):
                                table.rows[row_idx].cells[col_idx].text = cell_text
            
            # 分隔线
            elif re.match(r'^[\-=─]{3,}$', line):
                doc.add_paragraph('─' * 30)
            
            # 普通段落
            else:
                # 处理粗体和内联格式
                p = doc.add_paragraph()
                self._add_formatted_text(p, line)
            
            i += 1
        
        # 保存文档
        doc.save(output_path)
    
    def _add_formatted_text(self, paragraph, text):
        """
        在段落中添加格式化文本（支持 **粗体**）
        """
        # 简单的粗体处理
        parts = re.split(r'(\*\*.*?\*\*)', text)
        
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            else:
                paragraph.add_run(part)
    
    def _markdown_to_wechat(self, markdown_text: str) -> str:
        """
        将 Markdown 转换为微信友好的纯文本
        
        转换规则：
        - # 标题 → 📌【标题】（加粗用emoji）
        - ## 标题 → ▸ 标题
        - **粗体** → 【粗体】
        - - 列表 → · 列表
        - 表格 → 保留简单格式
        - 代码块 → 移除
        """
        lines = markdown_text.split('\n')
        result = []
        
        in_code_block = False
        
        for line in lines:
            # 跳过代码块
            if line.startswith('```'):
                in_code_block = not in_code_block
                continue
            if in_code_block:
                continue
            
            # 一级标题
            if line.startswith('# ') and not line.startswith('## '):
                text = line[2:].strip()
                result.append(f"\n━━━━━━━━━━━━━━━━")
                result.append(f"📌【{text}】")
                result.append("━━━━━━━━━━━━━━━━\n")
            
            # 二级标题
            elif line.startswith('## ') and not line.startswith('### '):
                text = line[3:].strip()
                result.append(f"\n▸ {text}")
            
            # 三级标题
            elif line.startswith('### '):
                text = line[4:].strip()
                result.append(f"\n» {text}")
            
            # 列表
            elif line.startswith('- ') or line.startswith('* '):
                text = line[2:].strip()
                # 转换粗体
                text = re.sub(r'\*\*(.*?)\*\*', r'【\1】', text)
                result.append(f"  · {text}")
            
            # 数字列表
            elif re.match(r'^\d+\.\s', line):
                text = re.sub(r'^\d+\.\s', '', line).strip()
                text = re.sub(r'\*\*(.*?)\*\*', r'【\1】', text)
                result.append(f"  {text}")
            
            # 分隔线
            elif re.match(r'^[\-=─]{3,}$', line):
                result.append("\n──────────────────────────────\n")
            
            # 表格行（保持原样，或简化）
            elif line.startswith('|'):
                # 跳过分隔行
                if not re.match(r'^\|[\s\-:]+\|', line):
                    result.append(line)
            
            # 普通文本
            else:
                # 转换粗体
                text = re.sub(r'\*\*(.*?)\*\*', r'【\1】', line)
                if text.strip():
                    result.append(text)
        
        return '\n'.join(result)

    async def _setup_knowledge_base(self):
        """
        挂载 PDF 知识库（赠金账户优化版）
        策略：尝试复用已上传的文件，失败时重新上传
        """
        pdf_path = "data/you_are_what_you_eat.pdf"
        if not os.path.exists(pdf_path) or not self.genai_client:
            return
        
        try:
            # 尝试从环境变量获取已上传的 file_ref（减少重复上传）
            cached_file_name = os.getenv("PDF_FILE_REF")
            
            if cached_file_name:
                try:
                    print(f"🔍 [System] 尝试复用已上传的 PDF: {cached_file_name}...", flush=True)
                    self.file_ref = self.genai_client.files.get(name=cached_file_name)
                    print(f"✅ [System] PDF 复用成功！无需重新上传。", flush=True)
                    return
                except Exception as reuse_error:
                    print(f"⚠️ [System] 无法复用（{reuse_error}），将重新上传...", flush=True)
            
            # 重新上传（赠金账户每次上传都计费，但无法使用 Cache API）
            print("📤 [System] 正在上传 PDF 知识库（约 4MB，需 10-30 秒）...", flush=True)
            self.file_ref = self.genai_client.files.upload(file=pdf_path)
            print(f"✅ [System] PDF 上传成功！ID: {self.file_ref.name}", flush=True)
            print(f"💡 [提示] 可设置环境变量以复用: export PDF_FILE_REF='{self.file_ref.name}'", flush=True)
            
        except Exception as e:
            print(f"💥 [System] PDF 挂载失败: {e}", flush=True)

    # ========== 关键 1：@ 消息入口（用户 @ agent 时触发）==========
    async def on_channel_mention(self, context: ChannelMessageContext):
        """
        处理 @ 消息（用户 @bc-intake 时触发）
        这是主要的消息处理入口！
        """
        print(f"🔔 [MENTION] on_channel_mention called, role={self.role_type}", flush=True)
        await self._process_channel_message(context)
    
    # ========== 关键 2：普通频道消息入口（不 @ 时触发）==========
    async def on_channel_post(self, context: ChannelMessageContext):
        """
        处理普通频道消息（不 @ agent 时触发）
        通常不会用到，因为我们要求用户 @ agent
        """
        print(f"🔔 [POST] on_channel_post called, role={self.role_type}", flush=True)
        # 普通消息也可以处理，但我们跳过
        # await self._process_channel_message(context)
        pass
    
    # ========== 核心消息处理逻辑 ==========
    async def _process_channel_message(self, context: ChannelMessageContext):
        """
        频道消息处理核心逻辑
        - intake：收集需求，输出结构化文档
        - content：基于 PDF + 膳食规则生成讲书内容
        - ops：生成可执行物料包
        """
        print(f"🔔 [PROCESS] _process_channel_message called, role={self.role_type}", flush=True)
        try:
            incoming = getattr(context, "incoming_event", None)
            print(f"🔔 [DEBUG] incoming={incoming is not None}", flush=True)
            
            payload = getattr(incoming, "payload", {}) or {}
            print(f"🔔 [DEBUG] payload={payload}", flush=True)
            
            user_text = (payload.get("content", {}) or {}).get("text", "").strip()
            print(f"🔔 [DEBUG] user_text length={len(user_text)}, text={user_text[:100] if user_text else 'EMPTY'}", flush=True)
            
            if not user_text:
                print(f"⚠️ [DEBUG] user_text is empty, returning", flush=True)
                return

            ws = self.workspace()
            channel = context.channel
            reply_to = incoming.id
            source_id = context.source_id

            print(f"🧭 [Channel] role={self.role_type} ch={channel} from={source_id}", flush=True)
            
            # 检查消息是否 @ 了当前 Agent（关键修复！）
            agent_mentions = [
                f"@bc-{self.role_type}",  # @bc-intake, @bc-content, @bc-ops
                f"@{self.role_type}",      # @intake, @content, @ops
            ]
            
            # 如果消息没有 @ 当前 Agent，忽略
            if not any(mention in user_text for mention in agent_mentions):
                print(f"⏭️  [Channel] 消息未 @ 当前 Agent，跳过", flush=True)
                return
            
            print(f"✅ [Channel] 消息 @ 了当前 Agent，开始处理", flush=True)

            # intake：收集需求，输出结构化文档
            if self.role_type == "intake":
                await ws.channel(channel).reply(reply_to, "✅【INTAKE】已收到。我正在整理需求...")

                intake_out = await self._execute_reasoning(user_text)
                
                # 自动保存到文件（三种格式）
                saved_path = self._save_output(intake_out)
                base_name = os.path.splitext(os.path.basename(saved_path))[0]
                
                guide = "\n\n" + "━" * 50 + "\n"
                guide += "💾 已生成多格式输出：\n"
                guide += f"  📄 Markdown: output/{base_name}.md\n"
                guide += f"  📘 Word文档: output/{base_name}.docx\n"
                guide += f"  📱 微信版: output/{base_name}_wechat.txt\n\n"
                guide += "📋 下一步：打开任一文件，复制内容，@bc-content 并粘贴。"
                
                await ws.channel(channel).reply(reply_to, f"🧾【INTAKE 输出】\n{intake_out}{guide}")
                return

            # content：基于 PDF + 膳食规则生成讲书内容（分天处理）
            if self.role_type == "content":
                rules_status = "✅ 膳食规则已加载" if self.rules_content else "⚠️ 膳食规则未加载"
                pdf_status = "✅ PDF 已挂载" if self.file_ref else "⚠️ PDF 未挂载"
                nutrition_status = "✅ 营养速查表已加载" if self.nutrition_content else "⚠️ 营养速查表未加载"
                
                # 从用户输入中提取天数
                import re
                days_match = re.search(r'(\d+)\s*天', user_text)
                total_days = int(days_match.group(1)) if days_match else 3
                
                await ws.channel(channel).reply(reply_to, f"""🧠【CONTENT】已接单，开始分天生成逐字稿。
{rules_status}
{pdf_status}
{nutrition_status}

📅 计划生成 **{total_days} 天**的讲书逐字稿
⏱️ 预计耗时：{total_days * 1} - {total_days * 2} 分钟
🔄 每天生成完成后会实时更新进度...""")

                # 第一步：生成主题大纲
                outline_prompt = f"""
{user_text}

请为这个 {total_days} 天的读书会生成【主题大纲】。

要求：
1. 每天一个主题，围绕《你是你吃出来的》的核心章节
2. 主题要有递进关系，从基础到深入
3. 最后一天要包含销讲环节

输出格式（只输出大纲，不要详细内容）：
Day 1：[主题名称] - [一句话描述]
Day 2：[主题名称] - [一句话描述]
...
Day {total_days}：[主题名称] - [一句话描述 + 销讲专场]
"""
                outline = await self._execute_reasoning(outline_prompt)
                await ws.channel(channel).reply(reply_to, f"📋 【大纲已生成】\n{outline}\n\n🔄 开始逐天生成详细逐字稿...")
                
                # 第二步：逐天生成内容
                all_content = [f"# 《你是你吃出来的》{total_days} 天读书会逐字稿\n\n{outline}\n\n---\n"]
                
                for day in range(1, total_days + 1):
                    # 确定第四部分的内容类型
                    if day == 1:
                        part4_type = "🌱 产品种草"
                        part4_desc = "轻描淡写，激发好奇，不要硬推"
                    elif day == total_days:
                        part4_type = "🎯 产品差异化 + 销讲"
                        part4_desc = "对比竞品，强调独特优势，完整销讲：痛点共情→科学解释→用户见证→产品介绍→促单→行动指令"
                    else:
                        part4_type = "💬 用户见证"
                        part4_desc = "真实案例，用户使用产品后的反馈和改变"
                    
                    day_prompt = f"""
{user_text}

【当前任务】生成 Day {day} 的完整逐字稿

【主题大纲参考】
{outline}

【输出结构 - 严格遵守】

# Day {day}：[从大纲中选择对应主题]

## 2.1 书中精华（15分钟逐字稿，约 1500 字）
- 像老师讲课一样，有开场白、过渡句
- 从《你是你吃出来的》PDF 中提取具体内容
- 引用书中原话，标注页码（如：PDF P22）
- 详细展开，不能只有大纲

## 2.2 延展知识（10分钟逐字稿，约 1000 字）
- 基于主理人专业背景的延伸讲解
- 补充书中没有但相关的营养知识
- 结合目标人群的实际场景

## 2.3 解决方案（5分钟逐字稿，约 500 字）
- 具体落地建议
- 使用营养速查表的精确数据（如：100g 鸡蛋含蛋白质 12.7g）
- 给出每日食谱建议

## {part4_type}（5-10分钟逐字稿，约 500 字）
{part4_desc}

【字数要求】
- 本天内容至少 3500 字
- 必须是可以直接朗读的【逐字稿】，不是大纲
- 每个小节都要有详细展开

【数据调用优先级】
- 涉及具体克数（g/ml）时 → 优先检索营养速查表
- 涉及医学逻辑/原理时 → 优先检索 PDF
- 涉及定量标准时 → 必须核对膳食指南（鸡蛋≤1个/天，盐<5g/天等）
"""
                    
                    await ws.channel(channel).reply(reply_to, f"⏳ 正在生成 Day {day}/{total_days}...")
                    day_content = await self._execute_reasoning(day_prompt)
                    all_content.append(day_content)
                    all_content.append("\n\n---\n\n")
                    
                    # 保存单天文件
                    self._save_output(day_content, suffix=f"day{day}")
                    await ws.channel(channel).reply(reply_to, f"✅ Day {day}/{total_days} 完成！（约 {len(day_content)} 字）")
                
                # 合并完整内容
                content_out = "\n".join(all_content)
                
                # 自动保存到文件（三种格式，内容可能很长！）
                saved_path = self._save_output(content_out)
                base_name = os.path.splitext(os.path.basename(saved_path))[0]
                
                guide = "\n\n" + "━" * 50 + "\n"
                guide += "💾 已生成多格式输出（内容较长）：\n"
                guide += f"  📄 Markdown: output/{base_name}.md\n"
                guide += f"  📘 Word文档: output/{base_name}.docx ← 可直接复制到微信公众号\n"
                guide += f"  📱 微信版: output/{base_name}_wechat.txt ← 朋友圈专用\n\n"
                guide += "📋 下一步：打开上述文件（推荐 Word），复制内容，@bc-ops 并粘贴。"
                
                await ws.channel(channel).reply(reply_to, f"📄【CONTENT 输出】\n{content_out}{guide}")
                return
            
            # ops：生成可执行物料包
            if self.role_type == "ops":
                rules_status = "✅ 膳食规则已加载" if self.rules_content else "⚠️ 膳食规则未加载"
                await ws.channel(channel).reply(reply_to, f"🧩【OPS】已接单。\n{rules_status}\n正在生成完整物料包（可能需要 1-2 分钟）...")

                # 强制要求输出完整物料包结构
                ops_prompt = f"""
{user_text}

【强制输出要求 - 必须严格遵守】

请输出完整的执行物料包，包含以下 5 个 Part：

# Part 3：时间轴与 SOP（约 1000 字）

## 3.1 总体时间线
| 阶段 | 时间 | 主要任务 |
|------|------|----------|
| 招募期 | D-7 至 D-1（7天） | ... |
| 交付期 | D1 至 D3（3天） | ... |
| 结营 | D3 下午 | ... |

## 3.2 招募期详细 SOP（7天）
| 日期 | 时间 | 动作 | 内容要点 |
|------|------|------|----------|
| D-7 | 08:00 | 朋友圈1 | ... |
...

## 3.3 交付期详细 SOP（3天）
...

# Part 4：招募期文案（21条，约 4000 字）

## 4.1 朋友圈文案（17条）

### D-7 第1条（预告）
---
[完整文案，可直接复制使用]
---

### D-7 第2条（痛点共鸣）
---
[完整文案]
---

... [继续写完 17 条]

## 4.2 公众号推文（2篇）
[完整标题+开头段落]

## 4.3 短视频脚本（7条）
[每条包含：画面描述+口播文案]

# Part 5：交付期文案（约 2000 字）

## 5.1 开营文案
[完整的欢迎语、群规、福利说明]

## 5.2 每日运营文案
### Day 1
- 早安问候：[完整文案]
- 作业引导：[完整文案]
- 晚安总结：[完整文案]

### Day 2
...

### Day 3
...

## 5.3 结营文案
[感谢语、成果回顾、销讲引导]

# Part 6：资源清单（约 500 字）

## 6.1 需要准备的图片素材
- [ ] 海报 x 3
- [ ] 产品图 x 5
...

## 6.2 需要准备的文档
...

## 6.3 时间投入估算
...

# Part 7：销讲资源包（约 1500 字）

## 7.1 销讲逐字稿
[完整的 10 分钟销讲脚本：痛点共情 → 科学解释 → 用户见证 → 产品介绍 → 促单 → 行动指令]

## 7.2 异议处理话术（5个）
| 常见异议 | 回应话术 |
|----------|----------|
| "太贵了" | ... |
...

## 7.3 接龙模板
[可直接复制的接龙格式]

## 7.4 成交后服务模板
[感谢语、使用指南、售后承诺]

【总字数要求】
- 总计至少 9000 字
- 每条文案必须是完整可用的，不是占位符

【内容要求】
- 文案风格：专业且温情
- 所有营养数据符合膳食指南
- 文案可直接复制使用，无需二次编辑
"""
                ops_out = await self._execute_reasoning(ops_prompt)
                
                # 自动保存到文件（三种格式）
                saved_path = self._save_output(ops_out)
                base_name = os.path.splitext(os.path.basename(saved_path))[0]
                
                final_guide = "\n\n" + "━" * 50 + "\n"
                final_guide += "💾 已生成多格式输出：\n"
                final_guide += f"  📄 Markdown: output/{base_name}.md\n"
                final_guide += f"  📘 Word文档: output/{base_name}.docx ← 直接在微信编辑器打开\n"
                final_guide += f"  📱 微信版: output/{base_name}_wechat.txt ← 逐条复制到朋友圈\n\n"
                final_guide += "✅ 全部完成！打开 output/ 文件夹，根据用途选择格式。\n"
                final_guide += "💡 使用建议：\n"
                final_guide += "  - 微信公众号：打开 .docx，直接复制到编辑器\n"
                final_guide += "  - 朋友圈文案：打开 _wechat.txt，逐条复制\n"
                final_guide += "  - 存档/修改：使用 .md 文件"
                
                await ws.channel(channel).reply(reply_to, f"📌【OPS 最终版 - 可直接使用的物料包】\n{ops_out}{final_guide}")
                return

        except Exception as e:
            print(f"💥 [Channel] 错误: {e}", flush=True)

    # ========== 推理（增强版：包含膳食规则约束） ==========
    async def _execute_reasoning(self, user_text: str) -> str:
        """
        执行 AI 推理
        - 自动注入膳食规则（如果已加载）
        - content 角色附带 PDF 知识库 + 营养速查表
        
        【数据调用优先级】
        - 涉及具体克数（g/ml）时 → 优先检索 nutrition_reference.md
        - 涉及医学逻辑时 → 优先检索 you_are_what_you_eat.pdf
        - 涉及定量标准时 → 必须核对 dietary_rules.md
        """
        if not self.genai_client:
            return "❌ API Key 缺失（GOOGLE_API_KEY）。"

        try:
            # 构建 prompt
            prompt_parts = [self.instruction]
            
            # 注入膳食规则（content 和 ops 角色）
            if self.rules_content and self.role_type in ("content", "ops"):
                rules_prompt = f"""
【重要约束 - 中国居民膳食指南2022】
以下是你必须严格遵守的膳食规则。在输出任何饮食建议时，必须符合这些规则：

{self.rules_content}

【约束提醒】
- 鸡蛋：每天最多1个，不弃蛋黄
- 食盐：每天<5g
- 牛奶：300-500ml/天
- 主食：每天≥150g
- 高血脂患者胆固醇：<200mg/天
"""
                prompt_parts.append(rules_prompt)
            
            # 注入营养速查表（content 角色）
            if self.nutrition_content and self.role_type == "content":
                nutrition_prompt = f"""
【食物营养速查表 - 用于精确数据引用】
当你需要引用具体的营养数据（如"100g 鸡蛋含蛋白质 12.7g"）时，请使用以下数据：

{self.nutrition_content}
"""
                prompt_parts.append(nutrition_prompt)
            
            prompt_parts.append(f"\n当前任务内容：{user_text}")
            prompt_content = "\n".join(prompt_parts)
            
            contents = [prompt_content]

            # content 角色可附带 PDF
            if self.role_type == "content" and self.file_ref:
                contents = [self.file_ref, prompt_content]

            # 根据角色设置不同的输出长度
            # content/ops 需要生成长文（~10000字），设置最大 tokens
            if self.role_type in ("content", "ops"):
                max_tokens = 8192  # 约 12000-15000 中文字符
            else:
                max_tokens = 2048  # intake 只需要简短输出
            
            resp = self.genai_client.models.generate_content(
                model=self.model_name,
                contents=contents,
                config={
                    "max_output_tokens": max_tokens,
                    "temperature": 0.7,  # 适度创意
                }
            )
            return resp.text if resp and getattr(resp, "text", None) else "⚠️ 无回复。"
        except Exception as e:
            return f"❌ 引擎报错: {str(e)}"
